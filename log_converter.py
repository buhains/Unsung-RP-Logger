import re
import codecs
import random
import json
from docx import Document
from docx.shared import RGBColor

def main():
    print("Hi Fern! Welcome to the Unsung RP Log converter.\nMake sure the log you want to convert is in the same folder as this program.\n")
    file_name = input("Please input the name of the file including the extension (e.g. log.txt):\n> ").strip().casefold()

    # Check file extension
    while True:
        try:
            file_title, file_ext = file_name.split(".")
        except:
            file_name = input("Sorry, please re-enter your file name with a .txt extension appended at the end:\n> ")
            continue
        else:
            if file_ext == "txt":
                break
            else:
                file_name = input("Sorry, only .txt files are accepted. Please input a file with the .txt extension:\n> ")
                file_title, file_ext = file_name.split(".")
                continue

    # File extension accepted
    print("Cleaning your log...")
    new_lines = clean_file(file_name,file_title)
    print(make_file(file_title, new_lines))

def clean_file(file_name, file_title):
    # Get file contents as
    f = codecs.open(file_name,"r",encoding="utf=8")
    raw_lines = f.readlines()
    new_lines = []

    # Remove tells, Gobchat artifacts, empty lines and command errors from log
    for line in raw_lines:
        artifacts = ["TellRecieve:", "TellSend:", "Error:", "Chatlogger", "Gobchat"]
        if not any(artifact in line for artifact in artifacts) and line != "\n" and line != "\r":
            new_lines.append(line)

    # Replacements
    replacements = {

        # Cross world linkshells
        "CrossWorldLinkShell_1:": "[OOC]:",
        "CrossWorldLinkShell_2:": "[DM]:",
        "CrossWorldLinkShell_3:": "[Player]:",

        # Chat channels
        " Party: ": " [Party]: ",
        " Alliance: ": " [Alliance]: ",
        " Emote: ": "",
        " Say:": ":",

        # Remove <p> and <br> lines
        "\n": "",
        "\r": "",

        # Remove party and alliance artefacts
        "": "",  # Party 1
        "": "",  # Party 2
        "": "",  # Party 3
        "": "",  # Party 4
        "": "",  # Party 5
        "": "",  # Party 6
        "": "",  # Party 7
        "": "",  # Party 8
        "": "",  # Alliance A
        "": "",  # Alliance B
        "": "",  # Alliance C

    }

    for original,replacement in replacements.items():
        new_lines = [line.replace(original,replacement) for line in new_lines]

    # Remove timestamps
    new_lines = [re.sub(r"\[\d\d\d\d-\d\d-\d\d\s\d\d:\d\d:[A-Za-z0-9]+(([+-]?(?=\.\d|\d)(?:\d+)?(?:\.?\d*))(?:[Ee]([+-]?\d+))?(:([+-]?(?=\.\d|\d)(?:\d+)?(?:\.?\d*))(?:[Ee]([+-]?\d+))?)+)\]\s+","",line) for line in new_lines]

    # Remove servers
    servers = [
        "Adamantoise",
        "Cactuar",
        "Faerie",
        "Gilgamesh",
        "Jenova",
        "Midgardsormr",
        "Sargatanas",
        "Siren",
        "Balmung",
        "Brynhildr",
        "Coeurl",
        "Diabolos",
        "Goblin",
        "Malboro",
        "Mateus",
        "Zalera",
        "Halicarnassus",
        "Maduin",
        "Marilith",
        "Seraph",
        "Behemoth",
        "Excalibur",
        "Exodus",
        "Famfrit",
        "Hyperion",
        "Lamia",
        "Leviathan",
        "Ultros"
    ]

    for server in servers:
        new_lines = [line.replace(f" [{server}]", "") for line in new_lines]

    return new_lines

def make_file(file_title, new_lines):

    # Create document
    document = Document()

    for line in new_lines:
        # Identify character name
        individualWords = line.split(" ")
        character_name = str(individualWords[0] + " " + individualWords[1]).strip()
        character_name = character_name.replace(":", "")

        # Identify rest of line
        restOfLine = line.replace(character_name,"")

        # Format character name
        p = document.add_paragraph("")
        character = p.add_run(character_name)
        rest = p.add_run(restOfLine)

        # Format character name with custom color
        r, g, b = color_name(character_name)
        character.font.color.rgb = RGBColor(r, g, b)
        character.bold = True

        # Make OOC text grey
        if "[OOC]:" in line:
            rest.font.color.rgb = RGBColor(141,141,141)

        # Make DM emote red and italicised
        elif "[DM]:" in line:
            character.font.color.rgb = RGBColor(225,0,0)
            character.italic = True
            rest.font.color.rgb = RGBColor(225,0,0)
            rest.italic = True

    # Save document
    new_file = str(f"{file_title}-clean.docx")
    document.save(new_file)

    return f"Success!\nYour cleaned log can be found in your current folder as: {new_file}"

def color_name(character_name):

    # Check if character already has assigned color
    f = open("name_colors.json")
    name_colors = json.load(f)
    f.close()

    # If character already has assigned color, retrieve
    for name in name_colors:
        if name.upper().strip() == character_name.upper().strip():
            r, g, b = name_colors[name].split("/")
            return int(r), int(g), int(b)

    # If character does not have assigned color, pick a random color
    with open("default_colors.json") as f:
        default_colors = json.load(f)
        rgb = random.choices(default_colors)[0]
        r, g, b = rgb.split("/")

    # Remove that random color from the pool of options
    default_colors.remove(rgb)
    default_colors = json.dumps(default_colors, indent=4)

    with open("default_colors.json", "w") as f:
        f.writelines(default_colors)

    # Add the character and their new color to the database
    name_colors.update({character_name: rgb})
    name_colors = json.dumps(name_colors, indent=4)

    with open("name_colors.json", "w") as f:
        f.writelines(name_colors)

    return int(r), int(g), int(b)

if __name__ == "__main__":
    main()
