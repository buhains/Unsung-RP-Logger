import re 
import codecs
from docx import Document
from docx.shared import RGBColor

def main():
    print("Hi Fern! Welcome to the Unsung RP Log converter.\nMake sure the log you want to convert is in the same folder as this program.\n")
    fileName = input("Please input the name of the file including the extension (e.g. log.txt):\n> ").strip()
    
    # Get file data
    fileTitle, fileExt = fileName.split(".")

    # Check file extension
    while True:
        if fileExt == "txt":
            break
        else:
            fileName = input("Sorry, only .txt files are accepted. Please input a file with the .txt extension:\n> ")
            fileTitle, fileExt = fileName.split(".")
            continue

    # File extension accepted
    print("Cleaning your log...")
    newLines = cleanFile(fileName,fileTitle)
    print(makeFile(fileTitle, newLines))

def cleanFile(fileName, fileTitle):
    # Get file contents as 
    f = codecs.open(fileName,"r",encoding="utf=8")
    rawLines = f.readlines()
    newLines = []
 
    # Remove tells, Gobchat artifacts, empty lines and command errors from log
    for line in rawLines:
        artifacts = ["TellRecieve:", "TellSend:", "Error:", "Chatlogger", "Gobchat"]
        if not any(artifact in line for artifact in artifacts) and line != "\n" and line != "\r":
            newLines.append(line)
    
    # Replacements
    replacements = {
        
        # Cross world linkshells
        "CrossWorldLinkShell_1:": "[OOC]:",
        "CrossWorldLinkShell_2:": "[DM]:",
        "CrossWorldLinkShell_3:": "[Player]:",

        # Chat channels
        " Party: ": " [Party]: ",
        " Alliance: ": " [Alliance]: ",
        " Emote: ": "",
        " Say:": ":",

        # Remove <p> and <br> lines
        "\n": "",
        "\r": ""
    }

    for original,replacement in replacements.items():
        newLines = [line.replace(original,replacement) for line in newLines]

    # Remove timestamps
    newLines = [re.sub(r"\[\d\d\d\d-\d\d-\d\d\s\d\d:\d\d:[A-Za-z0-9]+(([+-]?(?=\.\d|\d)(?:\d+)?(?:\.?\d*))(?:[Ee]([+-]?\d+))?(:([+-]?(?=\.\d|\d)(?:\d+)?(?:\.?\d*))(?:[Ee]([+-]?\d+))?)+)\]\s+","",line) for line in newLines]

    # Remove servers
    servers = [
        "Adamantoise",
        "Cactuar",
        "Faerie",
        "Gilgamesh",
        "Jenova",
        "Midgardsormr",
        "Sargatanas",
        "Siren",
        "Balmung",
        "Brynhildr",
        "Coeurl",
        "Diabolos",
        "Goblin",
        "Malboro",
        "Mateus",
        "Zalera",
        "Halicarnassus",
        "Maduin",
        "Marilith",
        "Seraph",
        "Behemoth",
        "Excalibur",
        "Exodus",
        "Famfrit",
        "Hyperion",
        "Lamia",
        "Leviathan",
        "Ultros"
    ]
    
    for server in servers:
        newLines = [line.replace(f" [{server}]", "") for line in newLines]
    
    return newLines

def makeFile(fileTitle, newLines):

    # Create document
    document = Document()

    for line in newLines:
        # Identify character name
        individualWords = line.split(" ")
        characterName = str(individualWords[0] + " " + individualWords[1]).strip()

        # Identify rest of line
        restOfLine = line.replace(characterName,"")

        # Add character name in bold
        p = document.add_paragraph("")
        p.add_run(characterName).bold = True
        rest = p.add_run(restOfLine)
        
        # Make OOC text grey
        if "[OOC]:" in line:
            rest.font.color.rgb = RGBColor(141,141,141)
        
        # Make DM emote red
        elif "[DM]:" in line:
            rest.font.color.rgb = RGBColor(225,0,0)

    # Save document
    newFile = str(f"{fileTitle}-clean.docx")
    document.save(newFile)

    return f"Success!\nYour cleaned log can be found in your current folder as: {newFile}"

if __name__ == "__main__":
    main()